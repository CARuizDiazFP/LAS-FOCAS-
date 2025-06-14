# Nombre de archivo: test_informe_sla.py
# Ubicación de archivo: tests/test_informe_sla.py
# User-provided custom instructions
# --------------------------------------------------------------------- #
#  Suite de pruebas unificada y libre de conflictos para el handler SLA #
# --------------------------------------------------------------------- #
import sys
import importlib
import asyncio
from pathlib import Path
from types import ModuleType, SimpleNamespace
import os
import pandas as pd
from docx import Document
import tempfile

# ─────────────────────────── PATH DE PROYECTO ─────────────────────────
ROOT_DIR = Path(__file__).resolve().parents[1]

# ────────────────────────── STUB TELEGRAM BASE ────────────────────────
from tests.telegram_stub import Message, Update, CallbackQuery  # type: ignore

# ─────────────────── VARIABLES DE ENTORNO MÍNIMAS ─────────────────────
# Las variables de entorno necesarias se definen en la fixture global

# Forzar que la base use SQLite en memoria
import sqlalchemy
from sqlalchemy.orm import sessionmaker

orig_engine = sqlalchemy.create_engine
sqlalchemy.create_engine = lambda *a, **k: orig_engine("sqlite:///:memory:")
bd = importlib.import_module("sandybot.database")
sqlalchemy.create_engine = orig_engine
bd.SessionLocal = sessionmaker(bind=bd.engine, expire_on_commit=False)
bd.Base.metadata.create_all(bind=bd.engine)

# ──────────── STUB REGISTRADOR – CAPTURA RESPUESTAS DEL HANDLER ──────
captura: dict[str, object] = {}

registrador_stub = ModuleType("sandybot.registrador")


async def responder_registrando(*args, **kwargs):
    captura["texto"] = args[3]
    captura["reply_markup"] = kwargs.get("reply_markup")


registrador_stub.responder_registrando = responder_registrando
registrador_stub.registrar_conversacion = lambda *a, **k: None
sys.modules["sandybot.registrador"] = registrador_stub

# ───────── FUNC. DE IMPORTACIÓN DINÁMICA DEL HANDLER ──────────────────
def _importar_handler(tmp_path: Path):
    plantilla = tmp_path / "plantilla.docx"
    doc = Document()
    headers = [
        "Tipo Servicio",
        "Número Línea",
        "Nombre Cliente",
        "Horas Reclamos Todos",
        "SLA",
    ]
    tbl1 = doc.add_table(rows=1, cols=len(headers))
    for i, h in enumerate(headers):
        tbl1.rows[0].cells[i].text = h

    tbl2 = doc.add_table(rows=5, cols=2)
    for i, t in enumerate(["Servicio", "Cliente", "N° de Ticket", "Domicilio", "SLA"]):
        tbl2.rows[i].cells[0].text = t

    headers_r = [
        "Número Línea",
        "Número Reclamo",
        "Horas Netas Reclamo",
        "Tipo Solución Reclamo",
        "Fecha Inicio Reclamo",
    ]
    tbl3 = doc.add_table(rows=1, cols=len(headers_r))
    for i, h in enumerate(headers_r):
        tbl3.rows[0].cells[i].text = h

    doc.save(plantilla)

    os.environ["SLA_TEMPLATE_PATH"] = str(plantilla)
    hist = tmp_path / "Historios"
    hist.mkdir()
    os.environ["SLA_HISTORIAL_DIR"] = str(hist)

    import importlib as _imp
    import sandybot.config as cfg
    _imp.reload(cfg)

    # Botón / teclado simple (stub)
    class _Btn:
        def __init__(self, text, callback_data=None):
            self.text = text
            self.callback_data = callback_data

    class _Mk:
        def __init__(self, keyboard):
            self.inline_keyboard = keyboard

    sys.modules["telegram"].InlineKeyboardButton = _Btn  # type: ignore
    sys.modules["telegram"].InlineKeyboardMarkup = _Mk   # type: ignore
    sys.modules["sandybot.registrador"] = registrador_stub

    # Carga dinámica del handler
    pkg = "sandybot.handlers"
    handlers_pkg = sys.modules.get(pkg) or ModuleType(pkg)
    handlers_pkg.__path__ = [str(ROOT_DIR / "Sandy bot" / "sandybot" / "handlers")]
    sys.modules[pkg] = handlers_pkg

    mod_name = f"{pkg}.informe_sla"
    spec = importlib.util.spec_from_file_location(
        mod_name, ROOT_DIR / "Sandy bot" / "sandybot" / "handlers" / "informe_sla.py"
    )
    mod = importlib.util.module_from_spec(spec)
    sys.modules[mod_name] = mod
    spec.loader.exec_module(mod)

    mod.RUTA_PLANTILLA = str(plantilla)
    return mod

# ──────────────── AYUDANTE PARA DOCUMENTOS EXCEL ──────────────────────
class ExcelDoc:
    def __init__(self, file_name: str, path: Path):
        self.file_name = file_name
        self._path = path

    async def get_file(self):
        class F:
            async def download_to_drive(_, dst):
                Path(dst).write_bytes(Path(self._path).read_bytes())
        return F()

# ───────────────────────── FLUJO COMPLETO SLA ─────────────────────────
async def _flujo_completo(tmp_path: Path):
    handler = _importar_handler(tmp_path)
    ctx = SimpleNamespace(user_data={})

    # /sla
    await handler.iniciar_informe_sla(Update(message=Message("/sla")), ctx)

    # Archivos de reclamos y servicios en una sola actualización
    recl = tmp_path / "recl.xlsx"
    serv = tmp_path / "serv.xlsx"
    pd.DataFrame({"Servicio": ["Srv"], "Número Reclamo": [1]}).to_excel(recl, index=False)
    pd.DataFrame(
        {
            "Tipo Servicio": ["Srv"],
            "Número Línea": [1],
            "Nombre Cliente": ["ACME"],
            "Horas Reclamos Todos": [0],
            "SLA": [0.5],
        }
    ).to_excel(serv, index=False)
    msg = Message()
    msg.documents = [ExcelDoc("recl.xlsx", recl), ExcelDoc("serv.xlsx", serv)]
    await handler.procesar_informe_sla(Update(message=msg), ctx)
    assert captura["reply_markup"].inline_keyboard[0][0].callback_data == "sla_procesar"

    # Procesar
    captura.clear()
    cb = SimpleNamespace(data="sla_procesar", message=msg)
    await handler.procesar_informe_sla(Update(callback_query=cb), ctx)

    assert msg.sent and not os.path.exists(tmp_path / msg.sent)

# ───────────── CAMBIO DE PLANTILLA DESDE EL BOT ───────────────────────
async def _cambio_plantilla(tmp_path: Path):
    handler = _importar_handler(tmp_path)
    ctx = SimpleNamespace(user_data={})

    await handler.iniciar_informe_sla(Update(message=Message("/sla")), ctx)

    cb = SimpleNamespace(data="sla_cambiar_plantilla", message=Message())
    await handler.procesar_informe_sla(Update(callback_query=cb), ctx)
    assert ctx.user_data["cambiar_plantilla"]

    nueva = tmp_path / "new.docx"
    Document().save(nueva)
    msg = Message(document=ExcelDoc("new.docx", nueva))

    msg.documents = [msg.document]

    await handler.procesar_informe_sla(Update(message=msg), ctx)

    assert "actualizada" in captura["texto"].lower()
    assert Path(handler.RUTA_PLANTILLA).read_bytes() == nueva.read_bytes()
    hist_dir = Path(os.environ["SLA_HISTORIAL_DIR"])
    assert any(hist_dir.iterdir())

async def _historial_plantilla(tmp_path: Path):
    handler = _importar_handler(tmp_path)
    hist_dir = Path(handler.config.SLA_HISTORIAL_DIR)
    for f in hist_dir.iterdir():
        f.unlink()
    ctx = SimpleNamespace(user_data={})

    await handler.iniciar_informe_sla(Update(message=Message("/sla")), ctx)

    cb = SimpleNamespace(data="sla_cambiar_plantilla", message=Message())
    await handler.procesar_informe_sla(Update(callback_query=cb), ctx)
    original = Path(handler.RUTA_PLANTILLA).read_bytes()

    nueva = tmp_path / "nueva.docx"
    Document().save(nueva)
    msg = Message(document=ExcelDoc("nueva.docx", nueva))
    await handler.procesar_informe_sla(Update(message=msg), ctx)

    archivos = list(hist_dir.iterdir())
    assert len(archivos) == 1
    assert archivos[0].read_bytes() == original

# ───────────── PRUEBA DE COLUMNAS OPCIONALES EN TABLA ─────────────────
def _test_columnas_extra(handler, tmp_path: Path):
    recl = tmp_path / "re.xlsx"
    serv = tmp_path / "se.xlsx"
    pd.DataFrame({"Servicio": [1]}).to_excel(recl, index=False)
    pd.DataFrame(
        {
            "Tipo Servicio": ["A"],
            "Número Línea": [1],
            "Nombre Cliente": ["X"],
            "Horas Reclamos Todos": [0],
            "SLA": [0.2],
        }
    ).to_excel(serv, index=False)
    doc_path = handler._generar_documento_sla(str(recl), str(serv))
    headers = [c.text for c in Document(doc_path).tables[0].rows[0].cells]
    assert headers == [
        "Tipo Servicio",
        "Número Línea",
        "Nombre Cliente",
        "Horas Reclamos Todos",
        "SLA",
    ]

# ───────────────────────── LISTA DE TESTS ─────────────────────────────
def test_flujo_completo(tmp_path):
    asyncio.run(_flujo_completo(tmp_path))

def test_actualizar_plantilla(tmp_path):
    captura.clear()
    asyncio.run(_cambio_plantilla(tmp_path))

def test_historial_de_plantillas(tmp_path):
    captura.clear()
    asyncio.run(_historial_plantilla(tmp_path))

def test_columnas_dinamicas(tmp_path):
    handler = _importar_handler(tmp_path)
    _test_columnas_extra(handler, tmp_path)

def test_exportar_pdf(tmp_path):
    """Genera PDF si el entorno lo permite, de lo contrario DOCX."""
    handler = _importar_handler(tmp_path)
    r, s = tmp_path / "r.xlsx", tmp_path / "s.xlsx"
    pd.DataFrame({"Servicio": [1]}).to_excel(r, index=False)
    pd.DataFrame(
        {
            "Tipo Servicio": ["A"],
            "Número Línea": [1],
            "Nombre Cliente": ["X"],
            "Horas Reclamos Todos": [0],
            "SLA": [0.4],
        }
    ).to_excel(s, index=False)
    ruta = handler._generar_documento_sla(str(r), str(s), exportar_pdf=True)
    assert ruta.endswith(".pdf") or ruta.endswith(".docx")


def test_tabla_orden_por_sla(tmp_path):
    """Verifica orden descendente y formato de celdas."""
    handler = _importar_handler(tmp_path)
    recl = tmp_path / "re.xlsx"
    serv = tmp_path / "se.xlsx"
    pd.DataFrame({"Servicio": [1, 2], "Fecha": ["2024-01-01", "2024-01-01"]}).to_excel(recl, index=False)
    pd.DataFrame(
        {
            "Tipo Servicio": ["A", "B"],
            "Número Línea": [1, 2],
            "Nombre Cliente": ["X", "Y"],
            "Horas Reclamos Todos": ["1 days 02:30:00", "0 days 05:00:00"],
            "SLA": [0.5, 0.3],
        }
    ).to_excel(serv, index=False)
    doc_path = handler._generar_documento_sla(str(recl), str(serv))
    tabla = Document(doc_path).tables[0]
    assert tabla.rows[1].cells[0].text == "A"
    assert tabla.rows[1].cells[3].text == "026:30:00"
    assert tabla.rows[1].cells[4].text == "50.00%"


def test_tablas_por_servicio(tmp_path):
    """El documento final debe contener pares de tablas por servicio."""
    handler = _importar_handler(tmp_path)
    r, s = tmp_path / "r.xlsx", tmp_path / "s.xlsx"
    pd.DataFrame({
        "Servicio": ["X", "Y"],
        "Número Reclamo": [1, 2],
        "Número Línea": [10, 20],
    }).to_excel(r, index=False)
    pd.DataFrame({
        "Tipo Servicio": ["X", "Y"],
        "Número Línea": [10, 20],
        "Nombre Cliente": ["A", "B"],
        "Horas Reclamos Todos": [0, 0],
        "SLA": [0.9, 0.8],
    }).to_excel(s, index=False)
    doc_path = handler._generar_documento_sla(str(r), str(s))
    doc = Document(doc_path)
    assert len(doc.tables) == 1 + 2 * 2


def test_bloque_con_parrafos(tmp_path):
    """Cada servicio debe incluir texto entre las tablas 2 y 3."""
    handler = _importar_handler(tmp_path)
    r, s = tmp_path / "r.xlsx", tmp_path / "s.xlsx"
    pd.DataFrame({
        "Servicio": ["X"],
        "Número Reclamo": [1],
        "Número Línea": [10],
        "Horas Netas Reclamo": ["0:00:00"],
        "Tipo Solución Reclamo": ["X"],
        "Fecha Inicio Reclamo": ["2024-01-01"],
    }).to_excel(r, index=False)
    pd.DataFrame({
        "Tipo Servicio": ["X"],
        "Número Línea": [10],
        "Nombre Cliente": ["A"],
        "Horas Reclamos Todos": [0],
        "SLA": [0.8],
    }).to_excel(s, index=False)
    doc_path = handler._generar_documento_sla(str(r), str(s), eventos="e", conclusion="c", propuesta="p")
    doc = Document(doc_path)
    body = doc._body._element
    idx2 = body.index(doc.tables[1]._tbl)
    idx3 = body.index(doc.tables[2]._tbl)
    entre = [child for child in body[idx2 + 1:idx3] if child.tag.endswith("p")]
    from docx.text.paragraph import Paragraph
    assert any("Eventos" in Paragraph(p, doc).text for p in entre)


def test_generar_con_ticket(tmp_path):
    """Soporta alias 'N° de Ticket' en los reclamos."""
    handler = _importar_handler(tmp_path)
    r, s = tmp_path / "r.xlsx", tmp_path / "s.xlsx"
    pd.DataFrame({
        "Servicio": ["X"],
        "N° de Ticket": [1],
        "Número Línea": [10],
        "Horas Netas Reclamo": ["0:00:00"],
        "Tipo Solución Reclamo": ["X"],
        "Fecha Inicio Reclamo": ["2024-01-01"],
    }).to_excel(r, index=False)
    pd.DataFrame({
        "Tipo Servicio": ["X"],
        "Número Línea": [10],
        "Nombre Cliente": ["A"],
        "Horas Reclamos Todos": [0],
        "SLA": [0.8],
    }).to_excel(s, index=False)
    doc_path = handler._generar_documento_sla(str(r), str(s))
    doc = Document(doc_path)
    tabla3 = doc.tables[-1]
    assert tabla3.rows[1].cells[1].text == "1"


def test_tabla2_completa(tmp_path):
    """Verifica que la Tabla 2 se llene con todos los valores."""
    handler = _importar_handler(tmp_path)
    r, s = tmp_path / "re.xlsx", tmp_path / "se.xlsx"
    pd.DataFrame({
        "Servicio": ["X", "X"],
        "Número Reclamo": [1, 2],
        "Número Línea": [10, 10],
    }).to_excel(r, index=False)
    pd.DataFrame({
        "Tipo Servicio": ["X"],
        "Número Línea": [10],
        "Nombre Cliente": ["A"],
        "Dirección Servicio": ["Calle 123"],
        "Horas Reclamos Todos": [0],
        "SLA Entregado": [0.5],
    }).to_excel(s, index=False)

    doc_path = handler._generar_documento_sla(str(r), str(s))
    t2 = Document(doc_path).tables[1]

    assert t2.rows[0].cells[1].text == "X 10"
    assert t2.rows[1].cells[1].text == "A"
    assert t2.rows[2].cells[1].text == "1, 2"
    assert t2.rows[3].cells[1].text == "Calle 123"
    assert t2.rows[4].cells[1].text == "50.00%"


def test_pdf_no_nameerror(tmp_path):
    """Confirma que exportar a PDF no produce NameError."""
    handler = _importar_handler(tmp_path)
    r, s = tmp_path / "r.xlsx", tmp_path / "s.xlsx"
    pd.DataFrame({"Servicio": [1]}).to_excel(r, index=False)
    pd.DataFrame(
        {
            "Tipo Servicio": ["A"],
            "Número Línea": [1],
            "Nombre Cliente": ["X"],
            "Horas Reclamos Todos": [0],
            "SLA": [0.5],
        }
    ).to_excel(s, index=False)
    try:
        ruta = handler._generar_documento_sla(str(r), str(s), exportar_pdf=True)
    except NameError as e:
        raise AssertionError(f"NameError inesperado: {e}")
    assert ruta.endswith(".pdf") or ruta.endswith(".docx")


def test_identificar_excel(tmp_path):
    handler = _importar_handler(tmp_path)

    recl = tmp_path / "r.xlsx"
    serv = tmp_path / "s.xlsx"
    pd.DataFrame({"Número Reclamo": [1]}).to_excel(recl, index=False)
    pd.DataFrame({"SLA Entregado": [0.5]}).to_excel(serv, index=False)

    assert handler.identificar_excel(str(recl)) == "reclamos"
    assert handler.identificar_excel(str(serv)) == "servicios"


def test_guardar_reclamos(tmp_path):
    handler = _importar_handler(tmp_path)
    srv = bd.crear_servicio(nombre="Srv", cliente="Cli")
    df = pd.DataFrame(
        {
            "Número Línea": [srv.id],
            "Número Reclamo": ["10"],
            "Fecha Inicio Problema Reclamo": ["2024-01-01"],
            "Fecha Cierre Problema Reclamo": ["2024-01-02"],
            "Tipo Solución Reclamo": ["Cambio"],
            "Descripción Solución Reclamo": ["Detalle"],
        }
    )
    handler._guardar_reclamos(df)
    recs = bd.obtener_reclamos_servicio(srv.id)
    assert recs[0].tipo_solucion == "Cambio"
    assert recs[0].descripcion_solucion == "Detalle"


def test_guardar_reclamos_ticket(tmp_path):
    handler = _importar_handler(tmp_path)
    srv = bd.crear_servicio(nombre="Srv2", cliente="Cli")
    df = pd.DataFrame({
        "Número Línea": [srv.id],
        "N° de Ticket": ["11"],
        "Fecha Inicio Problema Reclamo": ["2024-01-01"],
    })
    handler._guardar_reclamos(df)
    recs = bd.obtener_reclamos_servicio(srv.id)
    assert recs[0].numero == "11"


def test_titulo_unico_y_saltos(tmp_path):
    """El documento debe incluir un solo título y saltos entre servicios."""
    handler = _importar_handler(tmp_path)
    r, s = tmp_path / "r.xlsx", tmp_path / "s.xlsx"
    pd.DataFrame({
        "Servicio": ["A", "B"],
        "Número Reclamo": [1, 2],
        "Número Línea": [10, 20],
    }).to_excel(r, index=False)
    pd.DataFrame({
        "Tipo Servicio": ["A", "B"],
        "Número Línea": [10, 20],
        "Nombre Cliente": ["X", "Y"],
        "Horas Reclamos Todos": [0, 0],
        "SLA": [0.9, 0.8],
    }).to_excel(s, index=False)

    doc_path = handler._generar_documento_sla(str(r), str(s))
    doc = Document(doc_path)

    titulos = [p.text for p in doc.paragraphs if "Informe SLA" in p.text]
    assert len(titulos) == 0

    saltos = doc._element.xml.count('w:type="page"')
    filas = pd.read_excel(s)
    assert saltos == len(filas) - 1

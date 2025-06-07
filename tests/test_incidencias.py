import sys
import importlib
import asyncio
from types import ModuleType
from pathlib import Path
# Importar Document para crear el archivo de prueba
from docx import Document

# Preparar rutas para importar el paquete
ROOT_DIR = Path(__file__).resolve().parents[1]
sys.path.append(str(ROOT_DIR / "Sandy bot"))

# Stub del módulo telegram requerido por sandybot.utils
telegram_stub = ModuleType("telegram")
class Message:
    def __init__(self, text=""):
        self.text = text
class CallbackQuery:
    def __init__(self, message=None):
        self.message = message
class Update:
    def __init__(self, message=None, edited_message=None, callback_query=None):
        self.message = message
        self.edited_message = edited_message
        self.callback_query = callback_query
telegram_stub.Update = Update
telegram_stub.Message = Message
sys.modules.setdefault("telegram", telegram_stub)

# Stub de dotenv requerido por config
dotenv_stub = ModuleType("dotenv")
dotenv_stub.load_dotenv = lambda *a, **k: None
sys.modules.setdefault("dotenv", dotenv_stub)

# Stub de openai para evitar llamadas reales
openai_stub = ModuleType("openai")
class AsyncOpenAI:
    def __init__(self, api_key=None):
        self.chat = type("c", (), {"completions": type("comp", (), {"create": lambda *a, **k: None})()})()
openai_stub.AsyncOpenAI = AsyncOpenAI
sys.modules.setdefault("openai", openai_stub)

# Stub de jsonschema requerido por gpt_handler
jsonschema_stub = ModuleType("jsonschema")
jsonschema_stub.validate = lambda *a, **k: None
jsonschema_stub.ValidationError = type("ValidationError", (Exception,), {})
sys.modules.setdefault("jsonschema", jsonschema_stub)

# Variables de entorno mínimas para Config
import os
for var in ["TELEGRAM_TOKEN", "OPENAI_API_KEY", "NOTION_TOKEN", "NOTION_DATABASE_ID", "DB_USER", "DB_PASSWORD"]:
    os.environ.setdefault(var, "x")

# Importar módulos de SandyBot
config_mod = importlib.import_module("sandybot.config")
gpt_module = importlib.import_module("sandybot.gpt_handler")

# Reemplazar gpt por un stub para registrar el mensaje
class GPTStub(gpt_module.GPTHandler):
    async def consultar_gpt(self, mensaje: str, cache: bool = True) -> str:
        self.last_msg = mensaje
        return "ok"

gpt_module.gpt = GPTStub()

incidencias = importlib.reload(importlib.import_module("sandybot.incidencias"))


def test_procesar_incidencias_docx(tmp_path):
    doc_path = tmp_path / "incidencias.docx"
    doc = Document()
    doc.add_paragraph("Primera linea")
    doc.add_paragraph("Segunda linea")
    doc.save(doc_path)

    texto = incidencias.extraer_texto_doc(doc_path)
    respuesta = asyncio.run(incidencias.procesar_incidencias_docx(str(doc_path)))
    assert respuesta == "ok"
    assert incidencias.gpt.last_msg == texto

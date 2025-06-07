from docx import Document
from pathlib import Path
from .gpt_handler import gpt


def extraer_texto_doc(ruta: str) -> str:
    """Devuelve el texto de un archivo `.docx` o `.doc`."""
    ruta = str(ruta)
    if ruta.lower().endswith(".docx"):
        doc = Document(ruta)
        return "\n".join(p.text for p in doc.paragraphs if p.text)
    if ruta.lower().endswith(".doc"):
        try:
            import textract  # type: ignore
        except Exception as exc:  # pragma: no cover - solo se usa si hay .doc
            raise ValueError("Lectura de .doc no soportada: falta textract") from exc
        texto_bytes = textract.process(ruta)
        return texto_bytes.decode("utf-8", errors="ignore").replace("\r", "")
    raise ValueError("Extensión de archivo no soportada")

async def procesar_incidencias_docx(ruta: str) -> str:
    """Extrae el texto y lo envía a GPT."""
    texto = extraer_texto_doc(ruta)
    return await gpt.consultar_gpt(texto)


async def procesar_incidencias_archivos(rutas: list[str], contexto: str | None = None) -> str:
    """Combina texto de varios archivos y lo envía a GPT.

    Se aceptan paths a documentos ``.docx`` o ``.doc`` de texto plano. Si se
    indica ``contexto`` se adjunta al final del mensaje.
    """
    textos: list[str] = []

    for ruta in rutas:
        path = Path(ruta)
        if path.suffix.lower() == ".docx":
            doc = Document(path)
            textos.append("\n".join(p.text for p in doc.paragraphs if p.text))
        else:
            textos.append(path.read_text(encoding="utf-8"))

    if contexto:
        cpath = Path(contexto)
        if cpath.suffix.lower() == ".docx":
            doc = Document(cpath)
            textos.append("\n".join(p.text for p in doc.paragraphs if p.text))
        else:
            textos.append(cpath.read_text(encoding="utf-8"))

    mensaje = "\n".join(textos)
    return await gpt.consultar_gpt(mensaje)

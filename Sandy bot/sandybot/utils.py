# + Nombre de archivo: utils.py
# + Ubicación de archivo: Sandy bot/sandybot/utils.py
# User-provided custom instructions
"""
Funciones de utilidad comunes para el bot
"""

import json
import logging
import unicodedata
from datetime import datetime
from typing import Dict, Any, Optional
from pathlib import Path
from telegram import Update, Message
import re
from .config import config

logger = logging.getLogger(__name__)

def normalizar_texto(texto: str) -> str:
    """
    Normaliza un string para comparaciones (elimina acentos, mayúsculas, etc)
    """
    return unicodedata.normalize('NFKD', texto).encode('ascii', 'ignore').decode('ascii').lower()

def normalizar_camara(texto: str) -> str:
    """Normaliza nombres de cámara eliminando acentos y abreviaturas."""
    t = normalizar_texto(texto)

    # Equivalencias comunes en direcciones
    reemplazos = {
        r"\bcam\.\b": "camara",
        r"\bcam\b": "camara",
        r"\bav\.\b": "avenida",
        r"\bav\b": "avenida",
        r"\bgral\.\b": "general",
        r"\bgral\b": "general",
        r"\bcra\.?\b": "carrera",
    }
    for patron, reemplazo in reemplazos.items():
        t = re.sub(patron, reemplazo, t)

    # Eliminar puntuación que pueda afectar la comparación
    t = re.sub(r"[.,;:]", "", t)
    t = re.sub(r"\s+", " ", t)
    return t.strip()

def cargar_json(ruta: Path) -> Dict:
    """
    Carga un archivo JSON de forma segura
    """
    try:
        with open(ruta, 'r', encoding='utf-8') as f:
            return json.load(f)
    except FileNotFoundError:
        return {}
    except json.JSONDecodeError as e:
        logger.error(f"Error al decodificar JSON de {ruta}: {e}")
        return {}
    except Exception as e:
        logger.error(f"Error al cargar {ruta}: {e}")
        return {}

def guardar_json(datos: Dict, ruta: Path) -> bool:
    """
    Guarda datos en un archivo JSON de forma segura
    """
    try:
        ruta.parent.mkdir(parents=True, exist_ok=True)
        with open(ruta, 'w', encoding='utf-8') as f:
            json.dump(datos, f, ensure_ascii=False, indent=2)
        return True
    except Exception as e:
        logger.error(f"Error al guardar {ruta}: {e}")
        return False

def timestamp_log() -> str:
    """
    Genera un timestamp para logs
    """
    return datetime.now().strftime("%Y-%m-%d %H:%M:%S")

def obtener_mensaje(update: Update) -> Optional[Message]:
    """
    Devuelve el objeto ``Message`` de un ``Update``.

    Se revisan las distintas propiedades del ``Update`` para encontrar un
    mensaje válido. Si no se encuentra, retorna ``None``.
    """
    if update.message:
        return update.message
    if update.edited_message:
        return update.edited_message
    if update.callback_query and update.callback_query.message:
        return update.callback_query.message
    return None

def cargar_destinatarios(cliente: str) -> list[str]:
    """Obtiene los destinatarios asociados a ``cliente``."""
    from .database import obtener_cliente_por_nombre

    cli = obtener_cliente_por_nombre(cliente)
    return cli.destinatarios if cli and cli.destinatarios else []


def guardar_destinatarios(cliente: str, destinatarios: list[str]) -> bool:
    """Actualiza los correos del cliente indicado."""
    from .database import obtener_cliente_por_nombre, Cliente, SessionLocal

    with SessionLocal() as session:
        cli = obtener_cliente_por_nombre(cliente)
        if not cli:
            cli = Cliente(nombre=cliente, destinatarios=destinatarios)
            session.add(cli)
        else:
            cli.destinatarios = destinatarios
        session.commit()
        return True


def rellenar_tabla_sla(ruta_docx: str, datos: list[dict]) -> 'Document':
    """Rellena la primera tabla de ``ruta_docx`` con los valores de ``datos``.

    Se asume que la tabla de la plantilla posee cinco columnas en el orden:
    ``Tipo Servicio``, ``Número Línea``, ``Nombre Cliente``,
    ``Horas Reclamos Todos`` y ``SLA Entregado``.
    """
    from docx import Document
    import pandas as pd

    doc = Document(ruta_docx)
    if not doc.tables:
        raise ValueError("La plantilla debe incluir al menos una tabla")

    tabla = doc.tables[0]

    while len(tabla.rows) > 1:
        tabla._tbl.remove(tabla.rows[1]._tr)

    df = pd.DataFrame(datos, columns=[
        "Tipo Servicio",
        "Número Línea",
        "Nombre Cliente",
        "Horas Reclamos Todos",
        "SLA Entregado",
    ])

    for _, fila in df.iterrows():
        celdas = tabla.add_row().cells
        celdas[0].text = str(fila["Tipo Servicio"])
        celdas[1].text = str(fila["Número Línea"])
        celdas[2].text = str(fila["Nombre Cliente"])
        celdas[3].text = str(fila["Horas Reclamos Todos"])
        celdas[4].text = str(fila["SLA Entregado"])

    return doc


def incrementar_contador(clave: str, ruta: Path | None = None) -> int:
    """Obtiene el próximo número diario para ``clave``.

    Se almacena un contador por día en ``ruta`` o en ``config.ARCHIVO_CONTADOR``
    si no se especifica. Incrementa el valor y lo devuelve.
    """
    fecha = datetime.now().strftime("%d%m%Y")
    destino = ruta or config.ARCHIVO_CONTADOR
    data = cargar_json(destino)
    key = f"{clave}_{fecha}"
    numero = data.get(key, 0) + 1
    data[key] = numero
    guardar_json(data, destino)
    return numero

# Nombre de archivo: informe_sla.py
# Ubicación de archivo: Sandy bot/sandybot/handlers/informe_sla.py
# User-provided custom instructions: Siemple escribe en español y explica en detalles para que sirven las lineas modificadas, agregadas o quitadas.
"""Handler para generar informes de SLA."""

from __future__ import annotations

import copy
import logging
import locale
import os
import shutil
import tempfile
from pathlib import Path
from types import SimpleNamespace
from typing import Optional, Sequence

import pandas as pd
from docx import Document
from telegram import InlineKeyboardButton, InlineKeyboardMarkup, Update
from telegram.ext import ContextTypes

# ▸ Dependencias opcionales para exportar/modificar Word en Windows
try:  # pragma: no cover
    import win32com.client as win32               # type: ignore
    import pythoncom                              # type: ignore
except Exception:  # pragma: no cover
    win32 = None
    pythoncom = None

from sandybot.config import config
from ..utils import obtener_mensaje
from .estado import UserState
from ..registrador import responder_registrando, registrar_conversacion
from .. import database as bd

# Plantilla
RUTA_PLANTILLA = config.SLA_PLANTILLA_PATH
logger = logging.getLogger(__name__)

# ─────────────────────────────── UTILIDADES ────────────────────────────────
def _guardar_reclamos(df: pd.DataFrame) -> None:
    """Vuelca los reclamos del DataFrame a la BD si no existen."""
    col_ticket = next(
        (c for c in ("Número Reclamo", "N° de Ticket") if c in df.columns), None
    )
    col_servicio = next(
        (c for c in ("Servicio", "Número Línea", "Número Primer Servicio") if c in df.columns),
        None,
    )
    if not col_ticket or not col_servicio:
        return

    for _, fila in df.iterrows():
        sid = fila.get(col_servicio)
        ticket = fila.get(col_ticket)
        if pd.isna(sid) or pd.isna(ticket):
            continue
        try:
            sid_int = int(str(sid).replace(".0", ""))
        except ValueError:
            continue

        # Evitar duplicados
        if any(r.numero == str(ticket) for r in bd.obtener_reclamos_servicio(sid_int)):
            continue

        fecha_ini, fecha_fin = None, None
        if "Fecha Inicio Problema Reclamo" in df.columns:
            fecha_ini = pd.to_datetime(fila["Fecha Inicio Problema Reclamo"], errors="coerce")
        if "Fecha Cierre Problema Reclamo" in df.columns:
            fecha_fin = pd.to_datetime(fila["Fecha Cierre Problema Reclamo"], errors="coerce")

        bd.crear_reclamo(
            sid_int,
            str(ticket),
            fecha_inicio=fecha_ini,
            fecha_cierre=fecha_fin,
            tipo_solucion=fila.get("Tipo Solución Reclamo"),
            descripcion_solucion=fila.get("Descripción Solución Reclamo"),
        )


def identificar_excel(path: str) -> str:
    """Devuelve 'reclamos' o 'servicios' según las columnas del Excel."""
    columnas = set(pd.read_excel(path, nrows=0).columns.str.strip())
    if {"Número Reclamo", "N° de Ticket"} & columnas:
        return "reclamos"
    if {"SLA Entregado", "Número Primer Servicio"} & columnas:
        return "servicios"
    raise ValueError(f"No se pudo identificar el tipo de Excel: {Path(path).name}")


# ─────────────────────────────── INTERFAZ TG ───────────────────────────────
async def iniciar_informe_sla(update: Update, context: ContextTypes.DEFAULT_TYPE) -> None:
    """Activa el modo SLA y solicita los dos archivos."""
    mensaje = obtener_mensaje(update)
    if not mensaje:
        logger.warning("No se recibió mensaje en iniciar_informe_sla")
        return

    user_id = update.effective_user.id
    UserState.set_mode(user_id, "informe_sla")
    context.user_data.clear()
    context.user_data["archivos"] = [None, None]          # [reclamos, servicios]

    # Botón “Actualizar plantilla”
    try:
        kb = InlineKeyboardMarkup(
            [[InlineKeyboardButton("Actualizar plantilla", callback_data="sla_cambiar_plantilla")]]
        )
    except Exception:  # stubs de test
        btn = SimpleNamespace(text="Actualizar plantilla", callback_data="sla_cambiar_plantilla")
        kb = SimpleNamespace(inline_keyboard=[[btn]])

    await responder_registrando(
        mensaje,
        user_id,
        "informe_sla",
        "Enviá el Excel de **reclamos** y luego el de **servicios** para generar el informe.",
        "informe_sla",
        reply_markup=kb,
    )


async def procesar_informe_sla(
    update: Update,
    context: ContextTypes.DEFAULT_TYPE,
    *,
    exportar_pdf: bool = False,
) -> None:
    """Orquesta la carga de Excel y la generación del informe."""
    mensaje = obtener_mensaje(update)
    if not mensaje:
        logger.warning("No se recibió mensaje en procesar_informe_sla")
        return

    user_id = update.effective_user.id
    archivos: list[str | None] = context.user_data.setdefault("archivos", [None, None])

    # ─── 1) Callback «cambiar plantilla» ─────────────────────────────
    if update.callback_query and update.callback_query.data == "sla_cambiar_plantilla":
        context.user_data["cambiar_plantilla"] = True
        await update.callback_query.message.reply_text("Adjuntá la nueva plantilla .docx.")
        return

    # ─── 2) Recibir nueva plantilla ─────────────────────────────────
    if context.user_data.get("cambiar_plantilla"):
        if getattr(mensaje, "document", None):
            await actualizar_plantilla_sla(mensaje, context)
        else:
            await responder_registrando(
                mensaje, user_id, getattr(mensaje, "text", ""),
                "Adjuntá el .docx para actualizar la plantilla.", "informe_sla",
            )
        return

    # ─── 3) Callback «procesar informe» / «exportar PDF» ─────────────
    if update.callback_query and update.callback_query.data in {"sla_procesar", "sla_pdf"}:
        try:
            ruta_final = _generar_documento_sla(
                *archivos,
                exportar_pdf=exportar_pdf or update.callback_query.data == "sla_pdf",
            )
            with open(ruta_final, "rb") as f:
                await update.callback_query.message.reply_document(f, filename=Path(ruta_final).name)
            registrar_conversacion(
                user_id, "informe_sla", f"Documento {Path(ruta_final).name} enviado", "informe_sla"
            )
        except Exception as exc:  # pragma: no cover
            logger.error("Error generando SLA: %s", exc)
            await update.callback_query.message.reply_text("💥 Algo falló generando el informe de SLA.")
        finally:
            for p in archivos:
                if p:
                    Path(p).unlink(missing_ok=True)
            Path(ruta_final).unlink(missing_ok=True)
            context.user_data.clear()
            UserState.set_mode(user_id, "")
        return

    # ─── 4) Recepción de uno o varios Excel ─────────────────────────
    docs: list = [
        d for d in (getattr(mensaje, "document", None), *getattr(mensaje, "documents", [])) if d
    ]
    if docs:
        for doc in docs:
            tmp_fd, tmp_path = tempfile.mkstemp(suffix=".xlsx")
            os.close(tmp_fd)
            await (await doc.get_file()).download_to_drive(tmp_path)

            try:
                tipo = identificar_excel(tmp_path)
            except Exception as exc:  # pragma: no cover
                logger.warning("No se pudo clasificar %s: %s", doc.file_name, exc)
                tipo = "reclamos" if archivos[0] is None else "servicios"

            idx = 0 if tipo == "reclamos" else 1
            # Si ya hay uno de ese tipo, se sobre-escribe el antiguo
            if archivos[idx]:
                Path(archivos[idx]).unlink(missing_ok=True)
            archivos[idx] = tmp_path

        # ¿falta alguno?
        if None in archivos:
            faltante = "reclamos" if archivos[0] is None else "servicios"
            await responder_registrando(
                mensaje, user_id, docs[-1].file_name,
                f"Archivo guardado. Falta el Excel de {faltante}.", "informe_sla",
            )
            return

        # Botones Word / PDF
        try:
            kb = InlineKeyboardMarkup(
                [[InlineKeyboardButton("Procesar informe 🚀", callback_data="sla_procesar"),
                  InlineKeyboardButton("Exportar a PDF", callback_data="sla_pdf")]]
            )
        except Exception:
            p = SimpleNamespace(text="Procesar informe 🚀", callback_data="sla_procesar")
            q = SimpleNamespace(text="Exportar a PDF", callback_data="sla_pdf")
            kb = SimpleNamespace(inline_keyboard=[[p, q]])

        await responder_registrando(
            mensaje, user_id, docs[-1].file_name,
            "Archivos cargados. Elegí una opción.", "informe_sla",
            reply_markup=kb,
        )
        return

    # ─── 5) Sin adjuntos ni callbacks ────────────────────────────────
    await responder_registrando(
        mensaje, user_id, getattr(mensaje, "text", ""),
        "Adjuntá los Excel de reclamos y servicios para comenzar.", "informe_sla",
    )


# ──────────────────────── ACTUALIZAR PLANTILLA ────────────────────────
async def actualizar_plantilla_sla(mensaje, context):
    """Guarda la nueva plantilla y mueve la anterior a templates/Historios."""
    user_id = mensaje.from_user.id
    archivo = mensaje.document
    if not archivo.file_name.lower().endswith(".docx"):
        await responder_registrando(mensaje, user_id, archivo.file_name, "El archivo debe ser .docx.", "informe_sla")
        return

    try:
        f = await archivo.get_file()
        Path(RUTA_PLANTILLA).parent.mkdir(parents=True, exist_ok=True)

        if Path(RUTA_PLANTILLA).exists():
            ts = pd.Timestamp.now().strftime("%Y%m%d_%H%M%S")
            nombre_backup = f"{Path(RUTA_PLANTILLA).stem}_{ts}{Path(RUTA_PLANTILLA).suffix}"
            config.SLA_HISTORIAL_DIR.mkdir(parents=True, exist_ok=True)
            shutil.move(RUTA_PLANTILLA, config.SLA_HISTORIAL_DIR / nombre_backup)

        await f.download_to_drive(RUTA_PLANTILLA)
        texto = "Plantilla de SLA actualizada."
        context.user_data.pop("cambiar_plantilla", None)
    except Exception as exc:  # pragma: no cover
        logger.error("Error guardando plantilla SLA: %s", exc)
        texto = "No se pudo guardar la plantilla."

    await responder_registrando(mensaje, user_id, archivo.file_name, texto, "informe_sla")


# ───────────────────────── GENERADOR DE INFORME ─────────────────────────
def _generar_documento_sla(
    reclamos_xlsx: str,
    servicios_xlsx: str,
    *,
    eventos: str = "",
    conclusion: str = "",
    propuesta: str = "",
    exportar_pdf: bool = False,
) -> str:
    """Crea el informe SLA y devuelve la ruta del DOCX (o PDF)."""

    reclamos_df = pd.read_excel(reclamos_xlsx)
    reclamos_df.columns = reclamos_df.columns.str.replace(r"\s+", " ", regex=True).str.strip()

    servicios_df = pd.read_excel(servicios_xlsx)
    servicios_df.columns = servicios_df.columns.str.replace(r"\s+", " ", regex=True).str.strip()

    # Guarda reclamos en BD (ignora errores si BD no está configurada en tests)
    try:
        _guardar_reclamos(reclamos_df)
    except Exception:  # pragma: no cover
        logger.debug("No se pudo registrar reclamos en la BD (modo test)")

    if "SLA Entregado" in servicios_df.columns and "SLA" not in servicios_df.columns:
        servicios_df.rename(columns={"SLA Entregado": "SLA"}, inplace=True)

    # ── Normalizar fecha para el título ─────────────────────────────
    try:
        fecha = pd.to_datetime(reclamos_df.iloc[0].get("Fecha"))
        if pd.isna(fecha):
            raise ValueError
    except Exception:
        fecha = pd.Timestamp.today()

    for loc in ("es_ES.UTF-8", "es_ES", "es_AR.UTF-8", "es_AR"):
        try:
            locale.setlocale(locale.LC_TIME, loc)
            break
        except locale.Error:
            continue

    mes, anio = fecha.strftime("%B").upper(), fecha.strftime("%Y")

    # ── Cargar plantilla ────────────────────────────────────────────
    if not Path(RUTA_PLANTILLA).exists():
        raise ValueError(f"Plantilla de SLA no encontrada: {RUTA_PLANTILLA}")
    doc = Document(RUTA_PLANTILLA)
    cuerpo = doc._body._element

for p in list(doc.paragraphs):
    if "Informe SLA" in p.text:
        p.clear()

    nuevo_titulo = f"Informe SLA {mes} {anio}"
    if titulo_encontrado:
        titulo_encontrado.text = nuevo_titulo
    else:
        try:
            titulo = doc.add_heading(nuevo_titulo, level=0)
        except KeyError:  # pragma: no cover - compatibilidad con estilos
            titulo = doc.add_heading(nuevo_titulo, level=1)
        cuerpo.remove(titulo._p)
        cuerpo.insert(0, titulo._p)


    # ── Tabla principal (se asume que la plantilla contiene ≥1 tabla) ──
    if not doc.tables:
        raise ValueError("La plantilla debe incluir una tabla para el SLA")
    tabla_principal = doc.tables[0]

    # Copias de las tablas 2 y 3 para replicarlas por servicio
    tablas_plantilla = doc.tables[1:3]
    if len(tablas_plantilla) < 2:
        raise ValueError("La plantilla debe incluir tres tablas")
    tabla2_tpl, tabla3_tpl = [copy.deepcopy(t._tbl) for t in tablas_plantilla]
    cuerpo = doc._body._element

    # Párrafos entre las tablas 2 y 3 para replicar el bloque
    idx_t2 = cuerpo.index(doc.tables[1]._tbl)
    idx_t3 = cuerpo.index(doc.tables[2]._tbl)
    parrafos_tpl = []
    estilos_tpl = []
    from docx.text.paragraph import Paragraph
    for elem in list(cuerpo[idx_t2 + 1: idx_t3]):
        if elem.tag.endswith("p"):
            p = Paragraph(elem, doc)
            parrafos_tpl.append(p.text)
            estilos_tpl.append(p.style.name if p.style else None)
        cuerpo.remove(elem)

    if not parrafos_tpl:
        parrafos_tpl = [
            "Eventos sucedidos de mayor impacto en SLA:",
            "Conclusión:",
            "Propuesta de mejora:",
        ]
        estilos_tpl = [None] * len(parrafos_tpl)

    # Eliminar ejemplos originales de tablas
    cuerpo.remove(doc.tables[2]._tbl)
    cuerpo.remove(doc.tables[1]._tbl)

    # Borrar filas de ejemplo excepto encabezado
    while len(tabla_principal.rows) > 1:
        tabla_principal._tbl.remove(tabla_principal.rows[1]._tr)

    columnas_sla = [
        "Tipo Servicio",
        "Número Línea",
        "Nombre Cliente",
        "Horas Reclamos Todos",
        "SLA",
    ]
    faltantes = [c for c in columnas_sla if c not in servicios_df.columns]
    if faltantes:
        raise ValueError(f"Faltan columnas en servicios.xlsx: {', '.join(faltantes)}")

    # Normalizar SLA y formatear horas
    if "SLA" in servicios_df.columns:
        servicios_df["SLA"] = servicios_df["SLA"].apply(
            lambda v: float(str(v).replace(",", ".")) if not pd.isna(v) else 0
        )

    def _to_timedelta(val) -> pd.Timedelta:
        if pd.isna(val):
            return pd.Timedelta(0)
        try:
            return pd.to_timedelta(val)
        except Exception:
            return pd.to_timedelta(float(str(val).replace(",", ".")), unit="h")

    def _fmt_td(td: pd.Timedelta) -> str:
        total = int(td.total_seconds())
        h, m, s = total // 3600, (total % 3600) // 60, total % 60
        return f"{h:03d}:{m:02d}:{s:02d}"

    def _horas_decimal(val) -> str:
        """Convierte valores de horas a un número entero de horas."""
        if pd.isna(val) or val == "":
            return ""
        s = str(val).lower().replace(",", ".")
        s = s.replace("d\u00eda", "day").replace("d\u00edas", "day").replace("dias", "day")
        s = s.replace("horas", "hours").replace("hora", "hours")
        try:
            td = pd.to_timedelta(s)
            return str(int(td.total_seconds() // 3600))
        except Exception:
            try:
                return str(int(float(s)))
            except Exception:
                return s

    meses = ["ene", "feb", "mar", "abr", "may", "jun", "jul", "ago", "sep", "oct", "nov", "dic"]

    def _formatear_fecha(val) -> str:
        """Devuelve la fecha en formato DD-mes-YY en castellano."""
        try:
            fecha_v = pd.to_datetime(val)
        except Exception:
            return str(val)
        if pd.isna(fecha_v):
            return ""
        return f"{fecha_v.day:02d}-{meses[fecha_v.month - 1]}-{str(fecha_v.year)[2:]}"

    if "Horas Reclamos Todos" in servicios_df.columns:
        servicios_df["Horas Reclamos Todos"] = servicios_df["Horas Reclamos Todos"].apply(_to_timedelta).apply(_fmt_td)

    # Ordenar por SLA descendente y completar tabla principal
    servicios_ordenados = servicios_df.sort_values("SLA", ascending=False)
    for _, fila in servicios_ordenados[columnas_sla].iterrows():
        nueva = copy.deepcopy(tabla_principal.rows[0]._tr)
        tabla_principal._tbl.append(nueva)
        c = tabla_principal.rows[-1].cells
        c[0].text = str(fila["Tipo Servicio"])
        c[1].text = str(fila["Número Línea"])
        c[2].text = str(fila["Nombre Cliente"])
        c[3].text = str(fila["Horas Reclamos Todos"])
        c[4].text = f"{float(fila['SLA']) * 100:.2f}%"

    # ── Generar bloques por servicio ─────────────────────────────────
    col_ticket = next((c for c in ("Número Reclamo", "N° de Ticket") if c in reclamos_df.columns), None)
    col_match = "Número Línea" if "Número Línea" in reclamos_df.columns else None

    total_servicios = len(servicios_ordenados)
    for idx_srv, (_, srv) in enumerate(servicios_ordenados.iterrows()):
        # Tabla 2 con datos del servicio
        elem2 = copy.deepcopy(tabla2_tpl)
        cuerpo.append(elem2)
        t2 = doc.tables[-1]
        valores = {
            "servicio": f"{srv.get('Tipo Servicio', '')} {srv.get('Número Línea', '')}",
            "cliente": str(srv.get('Nombre Cliente', "")),
            "ticket": "",
            "domicilio": str(srv.get('Dirección Servicio', "")),
            "sla": str(srv.get('SLA', srv.get('SLA Entregado', "")))
        }
        if col_ticket and col_match:
            mask = reclamos_df[col_match] == srv.get(col_match)
            tickets = [str(t) for t in reclamos_df.loc[mask, col_ticket].dropna().unique()]
            valores["ticket"] = ", ".join(tickets)

        for row in t2.rows:
            titulo = row.cells[0].text.lower()
            if "servicio" in titulo:
                row.cells[1].text = valores["servicio"].strip()
            elif "cliente" in titulo:
                row.cells[1].text = valores["cliente"]
            elif "ticket" in titulo or "reclamo" in titulo:
                row.cells[1].text = valores["ticket"]
            elif "domicilio" in titulo:
                row.cells[1].text = valores["domicilio"]
            elif "sla" in titulo:
                row.cells[1].text = valores["sla"]

        # Párrafos informativos replicados desde la plantilla
        idx = cuerpo.index(elem2)
        for base, estilo in zip(parrafos_tpl, estilos_tpl):
            texto = base
            if "Eventos" in base:
                texto = f"Eventos sucedidos de mayor impacto en SLA: {eventos}".strip()
            elif "Conclusión" in base:
                texto = f"Conclusión: {conclusion}".strip()
            elif "Propuesta" in base:
                texto = f"Propuesta de mejora: {propuesta}".strip()
            p = doc.add_paragraph(texto, style=estilo)
            cuerpo.remove(p._p)
            cuerpo.insert(idx + 1, p._p)
            idx += 1

        # Tabla 3 con los reclamos del servicio
        elem3 = copy.deepcopy(tabla3_tpl)
        cuerpo.insert(idx + 1, elem3)
        t3 = doc.tables[-1]
        t3.style = "Table Grid"
        while len(t3.rows) > 1:
            t3._tbl.remove(t3.rows[1]._tr)
        if col_match:
            recls = reclamos_df[reclamos_df[col_match] == srv.get(col_match)]
        else:
            recls = reclamos_df
        total_h = 0.0
        for _, rec in recls.iterrows():
            cells = t3.add_row().cells
            cells[0].text = str(rec.get("Número Línea", ""))
            cells[1].text = str(rec.get(col_ticket, ""))
            horas = _horas_decimal(rec.get("Horas Netas Reclamo", ""))
            cells[2].text = horas
            try:
                total_h += float(horas)
            except Exception:
                pass
            cells[3].text = str(rec.get("Tipo Solución Reclamo", ""))
            cells[4].text = _formatear_fecha(rec.get("Fecha Inicio Reclamo", ""))

        fila_tot = t3.add_row().cells
        fila_tot[0].text = "Total"
        if total_h:
            fila_tot[2].text = str(int(total_h))

        # Salto de página entre servicios
        if idx_srv < total_servicios - 1:
            salto = doc.add_page_break()
            cuerpo.remove(salto._p)
            cuerpo.insert(cuerpo.index(elem3) + 1, salto._p)

    # ── Guardar DOCX temporal ───────────────────────────────────────
    fd, ruta_docx = tempfile.mkstemp(suffix=".docx")
    os.close(fd)
    doc.save(ruta_docx)

    # ── Modificación COM opcional (Windows) ─────────────────────────
    if win32 and pythoncom and os.name == "nt":
        _ajustar_titulo_com(ruta_docx, mes, anio)

    # ── Exportar PDF (si se solicitó) ───────────────────────────────
    if exportar_pdf:
        ruta_pdf = Path(ruta_docx).with_suffix(".pdf")
        convertido = False

        if win32 and os.name == "nt":
            try:
                word = win32.Dispatch("Word.Application")
                docx = word.Documents.Open(ruta_docx)
                docx.SaveAs(str(ruta_pdf), FileFormat=17)
                docx.Close()
                word.Quit()
                convertido = True
            except Exception:
                logger.warning("Conversión a PDF con win32 falló")

        if not convertido:
            try:
                from docx2pdf import convert  # type: ignore
                convert(ruta_docx, str(ruta_pdf))
                converted = True
            except Exception:
                logger.warning("Conversión a PDF con docx2pdf falló")

        if convertido:
            Path(ruta_docx).unlink(missing_ok=True)
            return str(ruta_pdf)

    return ruta_docx


# ──────────────────────────── AUXILIAR COM ───────────────────────────
def _ajustar_titulo_com(path: str, mes: str, anio: str) -> None:
    """Sobreescribe título mediante COM en Windows."""
    try:
        pythoncom.CoInitialize()
        word = win32.Dispatch("Word.Application")
        word.Visible = False
        doc = word.Documents.Open(path)
        titulo = f"Informe SLA {mes} {anio}"
        for shape in doc.Shapes:
            if shape.TextFrame.HasText and "Informe SLA" in shape.TextFrame.TextRange.Text:
                shape.TextFrame.TextRange.Text = titulo
        doc.Save()
        doc.Close()
        word.Quit()
    except Exception as exc:  # pragma: no cover
        logger.error("Error COM ajustando título SLA: %s", exc)
    finally:
        if pythoncom:
            pythoncom.CoUninitialize()
